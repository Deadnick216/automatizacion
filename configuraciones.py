from docx import Document
import os

# Crear un diccionario con las variables y sus valores personalizados
context = {
    'nombre': 'perin',
    'direccion': 'milagro',
}

# Leer el contenido del archivo de Word
doc = Document("cholo.docx")

# Función para reemplazar texto en los párrafos y mantener formato e imágenes
def reemplazar_texto_en_parrafos(parrafos, context):
    for p in parrafos:
        # Unir solo los runs con texto en una cadena temporal para facilitar el reemplazo
        texto_completo = ''.join([run.text for run in p.runs if run.text])
        
        # Realizar los reemplazos en la cadena completa de texto
        for key, value in context.items():
            texto_completo = texto_completo.replace(f'@{key}@', value)
        
        # Volver a colocar el texto modificado sin afectar runs vacíos (imágenes)
        i = 0
        for run in p.runs:
            if run.text:  # Solo reemplaza los runs con texto
                run.text = texto_completo[i:i + len(run.text)]
                i += len(run.text)

# Reemplazar en los párrafos del documento principal
reemplazar_texto_en_parrafos(doc.paragraphs, context)

# Reemplazar en los párrafos dentro de tablas, si existen
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            reemplazar_texto_en_parrafos(cell.paragraphs, context)

# Guardar el contenido modificado en un nuevo archivo
output_filename = "salida_rendered.docx"
doc.save(output_filename)

# Abrir el contrato modificado
os.system(f"start {output_filename}")
